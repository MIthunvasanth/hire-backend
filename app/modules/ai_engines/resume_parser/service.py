import json
import re
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

from docx import Document
from openai import AsyncOpenAI
from openai import OpenAIError
from pypdf import PdfReader

from app.core.config import settings
from app.utils.helpers import normalize_mongo_doc

from . import repository
from .schema import ResumeAtsRequest, ResumeAtsResponse, ResumeAtsResult, ResumeParseResponse, ResumeStoredResponse

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}


async def parse_resume_file(
    file_name: str,
    file_bytes: bytes,
    content_type: str | None,
    user_id: str | None = None,
    user_email: str | None = None,
) -> ResumeStoredResponse:
    extension = Path(file_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported file type. Please upload a PDF, DOC, DOCX, or TXT resume.")

    extracted_text = _extract_text(file_name=file_name, file_bytes=file_bytes)
    if not extracted_text.strip():
        raise ValueError("Could not extract readable text from the uploaded resume.")

    if not settings.openai_api_key:
        raise ValueError("OPENAI_API_KEY is not configured on the backend.")

    client = AsyncOpenAI(api_key=settings.openai_api_key)
    truncated_text = extracted_text[: settings.resume_parser_max_chars]

    try:
        response = await client.chat.completions.create(
            model=settings.openai_resume_model,
            response_format={"type": "json_object"},
            temperature=0.2,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You extract structured resume data for a career platform. "
                        "Return only valid JSON with this shape: "
                        "{"
                        '\"basics\": {\"name\": string|null, \"email\": string|null, \"phone\": string|null, \"location\": string|null, \"linkedin\": string|null, \"portfolio\": string|null}, '
                        '\"summary\": string|null, '
                        '\"skills\": [{\"name\": string, \"level\": string|null}], '
                        '\"experience\": [{\"title\": string, \"company\": string, \"duration\": string, \"location\": string|null, \"description\": string|null}], '
                        '\"education\": [{\"institution\": string, \"degree\": string, \"duration\": string|null, \"details\": string|null}], '
                        '\"projects\": [{\"name\": string, \"description\": string, \"technologies\": string[]}], '
                        '\"certifications\": [{\"name\": string, \"issuer\": string|null, \"year\": string|null}]'
                        "}. Use empty arrays when data is missing. Do not invent details."
                    ),
                },
                {
                    "role": "user",
                    "content": f"Extract structured resume data from this text:\n\n{truncated_text}",
                },
            ],
        )
    except OpenAIError as exc:
        raise ValueError(f"OpenAI parsing request failed: {exc}") from exc

    parsed_content = response.choices[0].message.content or "{}"
    parsed_json = _safe_parse_json(parsed_content)
    normalized_sections = _normalize_sections(parsed_json)

    parse_response = ResumeParseResponse(
        file_name=file_name,
        file_type=content_type or extension.lstrip("."),
        raw_text_excerpt=truncated_text[:1200],
        sections=normalized_sections,
    )

    created_at = datetime.now(timezone.utc).isoformat()
    record = {
        "file_name": parse_response.file_name,
        "file_type": parse_response.file_type,
        "raw_text_excerpt": parse_response.raw_text_excerpt,
        "sections": parse_response.sections.model_dump(),
        "user_id": user_id,
        "user_email": user_email.lower() if user_email else None,
        "created_at": created_at,
    }
    saved_id = await repository.create_parsed_resume(record)

    return ResumeStoredResponse(
        id=saved_id,
        created_at=created_at,
        user_id=user_id,
        user_email=user_email.lower() if user_email else None,
        **parse_response.model_dump(),
    )


async def get_latest_parsed_resume(user_id: str | None, user_email: str | None) -> ResumeStoredResponse | None:
    item = await repository.get_latest_parsed_resume(user_id=user_id, user_email=user_email)
    normalized = normalize_mongo_doc(item)
    if not normalized:
        return None
    return ResumeStoredResponse(**normalized)


async def get_parsed_resume_by_id(resume_id: str) -> ResumeStoredResponse | None:
    item = await repository.get_parsed_resume_by_id(resume_id)
    normalized = normalize_mongo_doc(item)
    if not normalized:
        return None
    return ResumeStoredResponse(**normalized)


async def generate_ats_resume(payload: ResumeAtsRequest) -> ResumeAtsResponse:
    resume_id = payload.resume_id
    if resume_id:
        resume_doc = await repository.get_parsed_resume_by_id(resume_id)
    else:
        resume_doc = await repository.get_latest_parsed_resume(user_id=payload.user_id, user_email=payload.user_email)

    resume = normalize_mongo_doc(resume_doc)
    if not resume:
        raise ValueError("No parsed resume found for ATS analysis")

    if not settings.openai_api_key:
        raise ValueError("OPENAI_API_KEY is not configured on backend")

    client = AsyncOpenAI(api_key=settings.openai_api_key)

    prompt = (
        "You are an expert ATS (Applicant Tracking System) resume optimization specialist and professional resume writer with 15+ years of experience helping candidates land interviews at top companies.\n\n"
        "You will receive structured data extracted from the user's existing resume including:\n"
        "- Personal details (name, contact, location)\n"
        "- Work experience (company, role, duration, responsibilities)\n"
        "- Education\n"
        "- Skills (technical + soft)\n"
        "- Projects\n"
        "- Certifications\n"
        "- Achievements\n\n"
        "Your job is to transform this raw extracted data into a fully ATS-optimized, professional resume output.\n\n"
        "## YOUR TASKS\n"
        "### 1. ATS OPTIMIZATION\n"
        "- Identify and inject relevant ATS keywords naturally into bullet points\n"
        "- Use standard section headings ATS systems recognize\n"
        "- Avoid tables, columns, graphics, headers/footers, and special characters that break ATS parsing\n"
        "- Ensure job title keywords match industry-standard terminology\n\n"
        "### 2. BULLET POINT TRANSFORMATION\n"
        "- Rewrite responsibilities into achievement-oriented bullet points\n"
        "- Use: action verb + task/skill + quantified result\n"
        "- If numbers are missing, suggest placeholders like [X%] or [N users]\n"
        "- Keep bullets concise and varied\n\n"
        "### 3. SKILLS SECTION\n"
        "- Separate skills into: Technical Skills | Tools & Technologies | Soft Skills\n"
        "- Add adjacent inferred skills marked [suggested]\n\n"
        "### 4. SUMMARY / OBJECTIVE SECTION\n"
        "- Write a 3-line professional summary\n"
        "- Include years/domain, top hard skills, and value proposition\n\n"
        "### 5. EDUCATION & CERTIFICATIONS\n"
        "- Keep formatting consistent\n"
        "- Surface role-critical certifications prominently\n\n"
        "### 6. PROJECTS\n"
        "- Format as Project | Tech Stack | Impact/Outcome\n"
        "- Add link placeholders if missing\n\n"
        "Return ONLY valid JSON with this exact structure: "
        "{"
        '"summary": "string", '
        '"contact": {"name": "string", "email": "string", "phone": "string", "location": "string", "linkedin": "string", "github": "string", "portfolio": "string"}, '
        '"experience": [{"company": "string", "role": "string", "duration": "string", "location": "string", "bullets": ["string"]}], '
        '"education": [{"degree": "string", "institution": "string", "year": "string", "gpa": "string", "relevant_coursework": ["string"]}], '
        '"skills": {"technical": ["string"], "tools": ["string"], "soft": ["string"], "suggested": ["string"]}, '
        '"projects": [{"name": "string", "tech_stack": ["string"], "description": "string", "link": "string"}], '
        '"certifications": ["string"], '
        '"ats_score": {"score": 0, "keyword_match": 0, "format_score": 0, "impact_score": 0, "missing_keywords": ["string"], "suggestions": ["string"]}'
        "}."
    )

    context = payload.target_context.model_dump()
    llm_input = {
        "resume": resume,
        "target_context": {
            "job_title": context.get("job_title") or "",
            "company_name": context.get("company_name") or "",
            "extracted_jd_keywords": context.get("extracted_jd_keywords") or [],
            "experience_level": context.get("experience_level") or "",
            "industry": context.get("industry") or "",
        },
    }

    try:
        response = await client.chat.completions.create(
            model=settings.openai_resume_model,
            response_format={"type": "json_object"},
            temperature=0.2,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": json.dumps(llm_input)},
            ],
        )
    except OpenAIError as exc:
        raise ValueError(f"ATS optimization request failed: {exc}") from exc

    parsed = _safe_parse_json(response.choices[0].message.content or "{}")
    parsed = _normalize_ats_result(parsed)
    ats_result = ResumeAtsResult(**parsed)

    created_at = datetime.now(timezone.utc).isoformat()
    record = {
        "resume_id": resume["id"],
        "user_id": payload.user_id or resume.get("user_id"),
        "user_email": (payload.user_email or resume.get("user_email") or "").lower() or None,
        "created_at": created_at,
        "target_context": payload.target_context.model_dump(),
        "result": ats_result.model_dump(),
    }
    saved_id = await repository.create_ats_result(record)

    return ResumeAtsResponse(
        id=saved_id,
        resume_id=resume["id"],
        user_id=record["user_id"],
        user_email=record["user_email"],
        created_at=created_at,
        target_context=payload.target_context,
        result=ats_result,
    )


async def get_latest_ats_for_user(user_id: str | None, user_email: str | None) -> ResumeAtsResponse | None:
    item = await repository.get_latest_ats_result(user_id=user_id, user_email=user_email)
    normalized = normalize_mongo_doc(item)
    if not normalized:
        return None
    return ResumeAtsResponse(**normalized)


def _safe_parse_json(content: str) -> dict:
    cleaned = content.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.strip("`")
        if cleaned.lower().startswith("json"):
            cleaned = cleaned[4:].strip()
    try:
        data = json.loads(cleaned)
        return data if isinstance(data, dict) else {}
    except json.JSONDecodeError as exc:
        raise ValueError("Model returned invalid JSON for resume parsing.") from exc


def _normalize_ats_result(data: dict) -> dict:
    if not isinstance(data, dict):
        data = {}

    def _str(value: object) -> str:
        return str(value).strip() if value is not None else ""

    def _list(value: object) -> list:
        return value if isinstance(value, list) else []

    ats_score = data.get("ats_score") if isinstance(data.get("ats_score"), dict) else {}

    def _score(name: str) -> int:
        raw = ats_score.get(name, 0)
        try:
            val = int(float(raw))
        except (TypeError, ValueError):
            val = 0
        return max(0, min(100, val))

    skills = data.get("skills") if isinstance(data.get("skills"), dict) else {}
    contact = data.get("contact") if isinstance(data.get("contact"), dict) else {}

    normalized = {
        "summary": _str(data.get("summary")),
        "contact": {
            "name": _str(contact.get("name")),
            "email": _str(contact.get("email")),
            "phone": _str(contact.get("phone")),
            "location": _str(contact.get("location")),
            "linkedin": _str(contact.get("linkedin")),
            "github": _str(contact.get("github")),
            "portfolio": _str(contact.get("portfolio")),
        },
        "experience": [],
        "education": [],
        "skills": {
            "technical": [_str(item) for item in _list(skills.get("technical")) if _str(item)],
            "tools": [_str(item) for item in _list(skills.get("tools")) if _str(item)],
            "soft": [_str(item) for item in _list(skills.get("soft")) if _str(item)],
            "suggested": [_str(item) for item in _list(skills.get("suggested")) if _str(item)],
        },
        "projects": [],
        "certifications": [_str(item) for item in _list(data.get("certifications")) if _str(item)],
        "ats_score": {
            "score": _score("score"),
            "keyword_match": _score("keyword_match"),
            "format_score": _score("format_score"),
            "impact_score": _score("impact_score"),
            "missing_keywords": [_str(item) for item in _list(ats_score.get("missing_keywords")) if _str(item)],
            "suggestions": [_str(item) for item in _list(ats_score.get("suggestions")) if _str(item)],
        },
    }

    for item in _list(data.get("experience")):
        if not isinstance(item, dict):
            continue
        normalized["experience"].append(
            {
                "company": _str(item.get("company")),
                "role": _str(item.get("role")),
                "duration": _str(item.get("duration")),
                "location": _str(item.get("location")),
                "bullets": [_str(bullet) for bullet in _list(item.get("bullets")) if _str(bullet)],
            }
        )

    for item in _list(data.get("education")):
        if not isinstance(item, dict):
            continue
        normalized["education"].append(
            {
                "degree": _str(item.get("degree")),
                "institution": _str(item.get("institution")),
                "year": _str(item.get("year")),
                "gpa": _str(item.get("gpa")),
                "relevant_coursework": [_str(course) for course in _list(item.get("relevant_coursework")) if _str(course)],
            }
        )

    for item in _list(data.get("projects")):
        if not isinstance(item, dict):
            continue
        normalized["projects"].append(
            {
                "name": _str(item.get("name")),
                "tech_stack": [_str(tech) for tech in _list(item.get("tech_stack")) if _str(tech)],
                "description": _str(item.get("description")),
                "link": _str(item.get("link")),
            }
        )

    return normalized


def _normalize_sections(data: dict) -> dict:
    def _optional_str(value: object) -> str | None:
        if value is None:
            return None
        text = str(value).strip()
        if not text or text.lower() in {"none", "null", "n/a", "na"}:
            return None
        return text

    basics = data.get("basics") if isinstance(data.get("basics"), dict) else {}
    summary = data.get("summary")

    def _list_of_dicts(key: str) -> list[dict]:
        value = data.get(key)
        if not isinstance(value, list):
            return []
        return [item for item in value if isinstance(item, dict)]

    skills = []
    for item in _list_of_dicts("skills"):
        name = str(item.get("name", "")).strip()
        if name:
            level_value = item.get("level")
            skills.append({"name": name, "level": _optional_str(level_value)})

    experience = []
    for item in _list_of_dicts("experience"):
        title = str(item.get("title", "")).strip()
        company = str(item.get("company", "")).strip()
        duration = str(item.get("duration", "")).strip()
        if title and company:
            experience.append(
                {
                    "title": title,
                    "company": company,
                    "duration": duration or "N/A",
                    "location": _optional_str(item.get("location")),
                    "description": _optional_str(item.get("description")),
                }
            )

    education = []
    for item in _list_of_dicts("education"):
        institution = str(item.get("institution", "")).strip()
        degree = str(item.get("degree", "")).strip()
        if institution and degree:
            education.append(
                {
                    "institution": institution,
                    "degree": degree,
                    "duration": _optional_str(item.get("duration")),
                    "details": _optional_str(item.get("details")),
                }
            )

    projects = []
    for item in _list_of_dicts("projects"):
        name = str(item.get("name", "")).strip()
        description = str(item.get("description", "")).strip()
        technologies_raw = item.get("technologies", [])
        technologies = [str(tech).strip() for tech in technologies_raw if str(tech).strip()] if isinstance(technologies_raw, list) else []
        if name and description:
            projects.append({"name": name, "description": description, "technologies": technologies})

    certifications = []
    for item in _list_of_dicts("certifications"):
        name = str(item.get("name", "")).strip()
        if name:
            certifications.append(
                {
                    "name": name,
                    "issuer": _optional_str(item.get("issuer")),
                    "year": _optional_str(item.get("year")),
                }
            )

    return {
        "basics": {
            "name": _optional_str(basics.get("name")),
            "email": _optional_str(basics.get("email")),
            "phone": _optional_str(basics.get("phone")),
            "location": _optional_str(basics.get("location")),
            "linkedin": _optional_str(basics.get("linkedin")),
            "portfolio": _optional_str(basics.get("portfolio")),
        },
        "summary": str(summary).strip() if isinstance(summary, str) and summary.strip() else None,
        "skills": skills,
        "experience": experience,
        "education": education,
        "projects": projects,
        "certifications": certifications,
    }


def _extract_text(file_name: str, file_bytes: bytes) -> str:
    extension = Path(file_name).suffix.lower()

    if extension == ".pdf":
        return _extract_pdf_text(file_bytes)
    if extension == ".doc":
        return _extract_doc_text(file_bytes)
    if extension == ".docx":
        return _extract_docx_text(file_bytes)
    if extension == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")

    raise ValueError("Unsupported file type.")


def _extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    table_cells.append(value)
    return "\n".join(paragraphs + table_cells)


def _extract_doc_text(file_bytes: bytes) -> str:
    decoded = file_bytes.decode("latin-1", errors="ignore")
    chunks = re.findall(r"[A-Za-z0-9@:/._,+#()' -]{4,}", decoded)
    cleaned_chunks = [chunk.strip() for chunk in chunks if any(character.isalpha() for character in chunk)]
    return "\n".join(cleaned_chunks)
